"""Generate Scoring Parameters DOCX with professional formatting."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn

doc = Document()

# Page setup — A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = section.bottom_margin = Cm(2.54)
section.left_margin = section.right_margin = Cm(3.18)
section.orientation = WD_ORIENTATION.PORTRAIT

# ── Style: Normal ──
sn = doc.styles["Normal"]
sn.font.name = "Calibri"
sn.font.size = Pt(11)
sn.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
sn.paragraph_format.line_spacing = 1.15
sn.paragraph_format.space_after = Pt(6)

# ── Style: Title ──
ts = doc.styles["Title"]
ts.font.name = "Calibri Light"
ts.font.size = Pt(28)
ts.font.bold = True
ts.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
ts.paragraph_format.space_after = Pt(4)

# ── Style: Subtitle ──
ss = doc.styles["Subtitle"]
ss.font.name = "Calibri"
ss.font.size = Pt(14)
ss.font.italic = True
ss.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
ss.paragraph_format.space_after = Pt(20)

# ── Style: Heading 1 ──
h1 = doc.styles["Heading 1"]
h1.font.name = "Calibri Light"
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(6)

# ── Style: List Bullet ──
lb = doc.styles["List Bullet"]
lb.font.name = "Calibri"
lb.font.size = Pt(11)
lb.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
lb.paragraph_format.space_after = Pt(4)
lb.paragraph_format.left_indent = Cm(1.27)

# ── Title Page ──
for _ in range(4):
    doc.add_paragraph()

t = doc.add_paragraph("Tracewise AI", style="Title")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

st = doc.add_paragraph("Hackathon Scoring Parameters", style="Subtitle")
st.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line.add_run("_" * 60)
run.font.color.rgb = RGBColor(0x43, 0x61, 0xEE)
run.font.size = Pt(10)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("AI Across the Product Development Lifecycle")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = meta2.add_run("Hackathon 2026  |  Intelligent Requirement-to-Test Traceability")
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ── Helper: add section ──
def add_section(heading, summary, bullets):
    doc.add_heading(heading, level=1)
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(8)
    for b in bullets:
        bp = doc.add_paragraph(b, style="List Bullet")

# ── Sections ──
add_section("1. Problem Definition",
    "Manual requirement-to-test traceability is slow (4+ hours per analysis), error-prone, "
    "and breaks when documents are scattered across formats. Coverage gaps go undetected until "
    "production failures. Our solution automates the entire requirement-to-test lifecycle in 30 seconds.",
    ["Multi-format upload: PDF, DOCX, Markdown, Text parsed automatically",
     "6 domain samples demonstrate real-world breadth (Software, EV, Civil, Mechanical, Payments, Aerospace)",
     "Backwards compatibility handles test cases written before requirements exist"])

add_section("2. Solution Architecture",
    "4-node LangGraph state machine: Requirement Refiner \u2192 Test Case Generator \u2192 "
    "Traceability Agent \u2192 Report Generator. Each node has independent fallback logic \u2014 "
    "no single point of failure. Express REST API with Prisma/SQLite persistence, "
    "Vite React frontend, Langfuse observability.",
    ["LangGraph StateGraph with typed annotation (graph.ts)",
     "Shared AgentState interface across all agents (state.ts)",
     "RESTful API: upload, analyze, results, report (routes/)",
     "Bidirectional pipeline: forward (requirement\u2192test) + reverse (orphan detection)"])

add_section("3. Engineering/Build Quality",
    "TypeScript strict mode throughout. Clean separation: agents/, routes/, parsers/. "
    "Every AI component has a non-AI fallback. Monorepo with npm workspaces. "
    "Professional CSS design system with custom properties and responsive breakpoints.",
    ["TypeScript strict mode enabled in tsconfig.json",
     "Fallback chain: LLM \u2192 regex, embeddings \u2192 word-level, Puppeteer \u2192 HTML",
     "895-line CSS design system with custom properties (index.css)",
     "Monorepo workspaces: backend + frontend"])

add_section("4. AI Appropriateness",
    "AI is used where it genuinely adds value \u2014 not everywhere. LLM handles "
    "unstructured\u2192structured extraction and test case generation. text-embedding-3-large "
    "generates semantic vectors for matching. Traditional code handles file parsing, "
    "report rendering, and database operations.",
    ["LLM called only in 2 of 10 source files (requirement-refiner.ts, test-case-generator.ts)",
     "text-embedding-3-large via Azure APIM generates 1536-dimension vectors for semantic matching",
     "File parsing uses pdf-parse/mammoth (no AI needed)",
     "Report generation is HTML\u2192PDF (no AI needed)"])

add_section("5. AI Implementation",
    "Structured prompt engineering with JSON output format. text-embedding-3-large generates "
    "1536-dimension vectors, matched via cosine similarity (threshold 0.6) with word-level "
    "Jaccard fallback (threshold 0.2). Bidirectional traceability: forward matching + "
    "orphan detection. Langfuse traces every LLM and embedding call.",
    ["4-level test generation prompt (Unit/Integration/System/Acceptance)",
     "text-embedding-3-large vectors \u2192 cosine similarity for requirement-test matching",
     "Orphan detection via reverse traceability pass",
     "Langfuse observability with trace/span logging on every LLM and embedding call"])

add_section("6. AI Impact",
    "Converts 4+ hours of manual traceability into 30 seconds. Bidirectional coverage "
    "analysis catches gaps humans miss. Multi-format compliance output (PDF/TXT/DOCX) "
    "for regulated industries. Graceful degradation ensures the system never blocks the user.",
    ["Semantic match: Google Login vs Verify OAuth sign-in \u2192 85% similarity (keyword: 15%)",
     "5 domain samples with ~42% coverage analysis completed in seconds",
     "Orphan detection finds test cases with no linked requirement (3 orphans in backwards-compat sample)",
     "Compliance documents generated in PDF, TXT, DOCX formats"])

add_section("7. Business Relevance",
    "PDLC traceability is a real pain point in regulated industries (automotive ISO 26262, "
    "banking PCI-DSS, civil ACI 318). Multi-domain support means one product serves many "
    "verticals. Compliance document generation is exactly what auditors require.",
    ["5 industry-specific sample datasets with domain-appropriate requirements",
     "Multi-format output (PDF/TXT/DOCX) for compliance delivery",
     "Backwards compatibility addresses real-world scenario: tests before specs finalized",
     "Cross-domain architecture reasoning documented in jury Q&A notes"])

add_section("8. Benefits & Viability",
    "Zero-dependency deployment: SQLite + local LLM calls. Works offline with regex fallback. "
    "Domain-agnostic architecture \u2014 swap the template, not the code. "
    "Graceful degradation at every layer ensures the system is always usable.",
    ["SQLite \u2014 no database server needed, zero-install",
     "LLM unavailable \u2192 regex parsing fallback (requirement-refiner.ts)",
     "Embeddings API fails \u2192 word-level similarity fallback (test-case-generator.ts)",
     "Puppeteer fails \u2192 HTML report fallback (report-generator.ts)"])

add_section("9. Working Demo",
    "Full end-to-end flow: upload \u2192 analyze \u2192 dashboard \u2192 PDF report. "
    "Visual dashboard with traceability matrix, coverage gaps, orphan detection, and "
    "type-filtered AI-generated test cases. 6 domain samples ready to test.",
    ["Drag-and-drop upload with file validation (Home.tsx)",
     "Dashboard: 5 summary cards, traceability matrix, gap analysis, orphan section, type filter",
     "6 domain subfolders + backwards-compat with test data",
     "Generated PDF reports from live analysis runs"])

add_section("10. Differentiation",
    "Bidirectional traceability (orphan detection) is uncommon \u2014 most tools only do "
    "forward mapping. Domain-agnostic engine with 5 industry samples. Backwards compatibility "
    "handles real-world incomplete data. Multi-format compliance output for regulated industries.",
    ["Reverse traceability pass detects test cases with no matching requirement",
     "6 domain subfolders: software, EV, civil, mechanical, payments, aerospace",
     "backwards-compat sample: 3 orphan test cases detected automatically",
     "Multi-format compliance output (PDF/TXT/DOCX) from markdown source"])

# Save
out = r"C:\Study\TraceWise_005\docs\Architecture\SCORING-PARAMETERS-FINAL.docx"
doc.save(out)
print(f"Saved: {out}")
