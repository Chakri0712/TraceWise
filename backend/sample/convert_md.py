"""Convert markdown files in this directory to .txt, .docx, and .pdf."""

import re
import os
from pathlib import Path

# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

def parse_md(text: str) -> list[dict]:
    """Parse markdown into a list of blocks: {type, level, text, items}."""
    blocks = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        # Heading
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue
        # Bullet list
        if re.match(r'^[-*]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i]):
                items.append(re.sub(r'^[-*]\s+', '', lines[i]).strip())
                i += 1
            blocks.append({"type": "bullets", "items": items})
            continue
        # Blank line
        if not line.strip():
            i += 1
            continue
        # Regular paragraph
        blocks.append({"type": "paragraph", "text": line.strip()})
        i += 1
    return blocks


def strip_md_inline(text: str) -> str:
    """Remove inline markdown formatting for plain text output."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text


# ---------------------------------------------------------------------------
# .txt generation
# ---------------------------------------------------------------------------

def to_txt(blocks: list[dict]) -> str:
    parts = []
    for b in blocks:
        if b["type"] == "heading":
            prefix = "#" * b["level"] + " "
            parts.append(prefix + strip_md_inline(b["text"]))
        elif b["type"] == "bullets":
            for item in b["items"]:
                parts.append("- " + strip_md_inline(item))
        elif b["type"] == "paragraph":
            parts.append(strip_md_inline(b["text"]))
    return "\n".join(parts) + "\n"


# ---------------------------------------------------------------------------
# .docx generation
# ---------------------------------------------------------------------------

def to_docx(blocks: list[dict], output_path: str, title: str):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.section import WD_ORIENTATION

    doc = Document()

    # Page setup — A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)
    section.orientation = WD_ORIENTATION.PORTRAIT

    # Style tuning
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    for n, size in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "Calibri Light"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)

    # Title
    doc.add_paragraph(title, style="Title")

    for b in blocks:
        if b["type"] == "heading":
            style = f"Heading {b['level']}" if b["level"] <= 3 else "Heading 3"
            doc.add_paragraph(b["text"], style=style)
        elif b["type"] == "bullets":
            for item in b["items"]:
                doc.add_paragraph(item, style="List Bullet")
        elif b["type"] == "paragraph":
            doc.add_paragraph(b["text"], style="Normal")

    doc.save(output_path)


# ---------------------------------------------------------------------------
# .pdf generation
# ---------------------------------------------------------------------------

def to_pdf(blocks: list[dict], output_path: str, doc_title: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"],
        fontSize=22, spaceAfter=20, textColor=colors.HexColor("#1F3A5F"),
    )
    h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=16, spaceBefore=14, spaceAfter=6,
        textColor=colors.HexColor("#1F3A5F"),
    )
    h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=13, spaceBefore=10, spaceAfter=4,
        textColor=colors.HexColor("#1F3A5F"),
    )
    h3 = ParagraphStyle(
        "H3", parent=styles["Heading3"],
        fontSize=11, spaceBefore=8, spaceAfter=4,
        textColor=colors.HexColor("#1F3A5F"),
    )
    body = ParagraphStyle(
        "Body", parent=styles["BodyText"],
        fontSize=10, leading=14, spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=body,
        leftIndent=18, bulletIndent=6, spaceAfter=3,
    )

    heading_styles = {1: h1, 2: h2, 3: h3}

    def chrome(canv, doc):
        canv.saveState()
        canv.setFont("Helvetica", 8)
        canv.setFillGray(0.4)
        canv.drawString(2.0 * cm, 1.5 * cm, doc_title)
        canv.drawRightString(A4[0] - 2.0 * cm, 1.5 * cm, f"Page {doc.page}")
        canv.restoreState()

    pdf = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=3.0 * cm, rightMargin=3.0 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=doc_title, author="Auto-generated",
    )

    story = []
    for b in blocks:
        if b["type"] == "heading":
            story.append(Paragraph(
                escape_xml(b["text"]),
                heading_styles.get(b["level"], h3),
            ))
        elif b["type"] == "bullets":
            for item in b["items"]:
                story.append(Paragraph(
                    f"<bullet>&bull;</bullet>{escape_xml(item)}",
                    bullet_style,
                ))
        elif b["type"] == "paragraph":
            story.append(Paragraph(escape_xml(b["text"]), body))

    pdf.build(story, onFirstPage=chrome, onLaterPages=chrome)


def escape_xml(text: str) -> str:
    """Escape XML special characters for reportlab Paragraph."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    directory = Path(__file__).parent

    md_files = sorted(directory.glob("*.md"))
    if not md_files:
        print("No .md files found.")
        return

    for md_path in md_files:
        stem = md_path.stem
        title = stem.replace("_", " ").replace("-", " ").title()
        text = md_path.read_text(encoding="utf-8")
        blocks = parse_md(text)

        # .txt
        txt_path = md_path.with_suffix(".txt")
        txt_path.write_text(to_txt(blocks), encoding="utf-8")
        print(f"Created {txt_path.name}")

        # .docx
        docx_path = md_path.with_suffix(".docx")
        to_docx(blocks, str(docx_path), title)
        print(f"Created {docx_path.name}")

        # .pdf
        pdf_path = md_path.with_suffix(".pdf")
        to_pdf(blocks, str(pdf_path), title)
        print(f"Created {pdf_path.name}")

    print("\nDone! All files generated.")


if __name__ == "__main__":
    main()
